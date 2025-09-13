"""
DOCX Document Processor Implementation.
"""
from typing import List, Dict, Any, Optional
import logging
import re

# Import python-docx for DOCX processing
try:
    import docx
except ImportError:
    docx = None

from embedding_vectorstore.document_processor import DocumentProcessor, TextChunker

logger = logging.getLogger(__name__)

class DocxProcessor(DocumentProcessor):
    """DOCX document processor"""
    
    def __init__(self):
        if docx is None:
            logger.warning("python-docx is not installed. DOCX processing will not be available.")
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if processor supports given file extension"""
        return file_extension.lower() in ['.docx'] and docx is not None
    
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None, 
                    chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """Process DOCX document and return list of text chunks"""
        if docx is None:
            raise ImportError("python-docx is required for DOCX processing")
            
        # Extract text
        text = await self._extract_text(file_path)
        
        # Extract metadata if not provided
        if metadata is None:
            metadata = {}
        doc_metadata = await self._extract_metadata(file_path)
        metadata.update(doc_metadata)
        
        # Clean text
        text = self._clean_text(text)
        
        # Chunk text
        chunker = TextChunker()
        chunks = chunker.chunk(text, chunk_size, chunk_overlap)
        
        return chunks
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Handle tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    
    async def _extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        doc = docx.Document(file_path)
        metadata = {}
        
        properties = doc.core_properties
        if properties:
            if properties.author:
                metadata['author'] = properties.author
            if properties.created:
                metadata['created'] = properties.created.isoformat()
            if properties.modified:
                metadata['modified'] = properties.modified.isoformat()
            if properties.title:
                metadata['title'] = properties.title
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()